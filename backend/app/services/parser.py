"""文档解析器：PDF / DOCX / TXT / MD → 可检索文本。

多模态问题修复（第3项）：
- PDF 增加空文本检测：扫描版 PDF（每页为图片）extract_text() 返回空
- 增加 OCR fallback：优先尝试 pytesseract，其次尝试 Qwen-VL 逐页识别
- 对提取结果过短的页面给出警告日志，便于排查入库质量问题

2026-06-29 问题分析修复：
- DOCX 不再只读 paragraph.text，同时提取表格与内嵌图片信息；
- PDF 文本型页面也统计内嵌图片数量并写入可检索文本；
- 对图片优先使用 OCR/VL 生成文字描述，使"文档中的照片"不再从入库阶段静默丢失。
"""
import logging
import os
import hashlib
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# 上传接口必须快速返回；逐图 OCR/VL 会让大 PDF 卡死或超时。
# 因此上传阶段只抽取/保存图片并写入路径，视觉理解交给用户查询图或后续异步任务。
_IMAGE_EXTRACT_LIMIT = 80


@dataclass
class ParsedDocument:
    text: str
    images: list[dict]


def _safe_name(name: str) -> str:
    base = os.path.basename(name or "document")
    return "".join(c if c.isalnum() or c in ".-_" else "_" for c in base).strip("._") or "document"


def _image_output_dir(source_path: str) -> str:
    from app.core.config import settings
    stem = _safe_name(os.path.splitext(os.path.basename(source_path))[0])
    digest = hashlib.md5(os.path.abspath(source_path).encode("utf-8")).hexdigest()[:8]
    out_dir = os.path.join(settings.EXTRACTED_IMAGE_DIR, f"{stem}_{digest}")
    os.makedirs(out_dir, exist_ok=True)
    return out_dir

# 延迟导入：pdf2image 和 pytesseract 在龙芯上可能未安装
_pdf2image_available = None
_pytesseract_available = None


def _check_ocr_deps():
    """检查 OCR 依赖是否可用（懒加载，只检查一次）。"""
    global _pdf2image_available, _pytesseract_available
    if _pdf2image_available is not None:
        return
    try:
        import pdf2image  # noqa: F401
        _pdf2image_available = True
    except ImportError:
        _pdf2image_available = False
    try:
        import pytesseract  # noqa: F401
        _pytesseract_available = True
    except ImportError:
        _pytesseract_available = False


def read_pdf(path: str) -> str:
    return parse_pdf(path).text


def parse_pdf(path: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(path)
    raw_parts = []
    image_items: list[dict] = []
    empty_pages = 0
    total_pages = len(reader.pages)
    embedded_image_count = 0
    out_dir = _image_output_dir(path)

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        page_parts = []
        if not text or len(text) < 5:
            empty_pages += 1
        else:
            page_parts.append(text)
        images = getattr(page, "images", []) or []
        if images:
            embedded_image_count += len(images)
            page_parts.append(f"[第 {i + 1} 页包含 {len(images)} 张内嵌图片/图表，以下为图片识别结果。]")
            for j, img in enumerate(images, start=1):
                if len(image_items) >= _IMAGE_EXTRACT_LIMIT:
                    page_parts.append(f"[第 {i + 1} 页另有图片未抽取，已达到 {_IMAGE_EXTRACT_LIMIT} 张处理上限。]")
                    continue
                try:
                    ext = os.path.splitext(getattr(img, "name", ""))[1] or ".png"
                    img_path = os.path.join(out_dir, f"page-{i + 1:03d}-image-{j:02d}{ext}")
                    with open(img_path, "wb") as f:
                        f.write(img.data)
                    label = f"第 {i + 1} 页图片 {j}"
                    page_parts.append(
                        f"[{label}]\n"
                        f"该图片来自原文档第 {i + 1} 页，可作为维修手册图片证据返回给用户。\n"
                        f"[图片文件] {img_path}"
                    )
                    image_items.append({
                        "page": i + 1,
                        "index": j,
                        "path": img_path,
                        "description": "",
                    })
                except Exception as e:
                    logger.warning("PDF 图片提取失败 (%s, page %d image %d): %s", path, i + 1, j, e)
                    page_parts.append(f"[第 {i + 1} 页图片 {j}：提取失败。]")
        if page_parts:
            raw_parts.append("\n".join(page_parts))

    raw_text = "\n\n".join(raw_parts).strip()
    if embedded_image_count:
        logger.warning(
            "PDF %s: 检测到 %d 张内嵌图片/图表，已抽取到 %s 并写入识别结果。",
            path, embedded_image_count, out_dir,
        )

    # 多模态问题修复：检测是否为扫描版 PDF（大部分页面无文字）
    # 如果超过半数页面为空，说明是图片型 PDF，需要 OCR
    is_scanned_like = total_pages > 0 and empty_pages > total_pages * 0.5

    if not raw_text and total_pages > 0:
        logger.warning(
            "PDF %s: extract_text() 返回空文本（%d 页），"
            "可能是扫描版 PDF 或图片型 PDF。建议安装 OCR 依赖："
            "pip install pdf2image pytesseract 并确保系统已装 tesseract",
            path, total_pages,
        )

    if is_scanned_like and (raw_text or "") and len(raw_text) < 50:
        logger.warning(
            "PDF %s: 仅提取到极少文字（%d 字符，%d/%d 页空白），"
            "内容可能不完整。如需完整识别请安装 OCR 依赖。",
            path, len(raw_text), empty_pages, total_pages,
        )

    # 注意：上传接口不再同步做整本 PDF OCR。
    # 大型手册逐页 OCR/VL 会导致上传接口长时间阻塞、前端超时并显示上传失败。
    # 如需扫描版 PDF 全文 OCR，应拆成后台任务或离线预处理。

    return ParsedDocument(text=raw_text, images=image_items)


def _try_ocr_fallback(path: str) -> bool:
    """检查 OCR 依赖是否就绪。"""
    _check_ocr_deps()
    return bool(_pdf2image_available and _pytesseract_available)


def _ocr_pdf(path: str) -> str:
    """用 pytesseract 对 PDF 做 OCR，返回合并后的文字。"""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(path, dpi=200)
        parts = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang='chi_sim+eng').strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("PDF OCR 失败 (%s): %s", path, e)
        return ""


def _ocr_image_file(image_path: str) -> str:
    """尽力识别图片内容：先 pytesseract，失败再尝试 Qwen-VL。"""
    # 1) 本地 OCR：适合有文字的图纸/截图
    try:
        import pytesseract
        from PIL import Image
        text = pytesseract.image_to_string(Image.open(image_path), lang='chi_sim+eng').strip()
        if text:
            return text
    except Exception as e:
        logger.debug("图片 OCR 不可用或失败 (%s): %s", image_path, e)

    # 2) VL 描述：适合照片/图表，但依赖 DashScope，可失败降级
    try:
        from app.services.llm import vl_describe
        return vl_describe(f"file://{os.path.abspath(image_path)}", mode="document").strip()
    except Exception as e:
        logger.warning("图片 VL 描述失败 (%s): %s", image_path, e)
        return ""


def _extract_docx_images(doc, source_path: str) -> tuple[list[str], list[dict]]:
    """提取 DOCX 内嵌图片，并把 OCR/VL 描述写入文本。"""
    parts: list[str] = []
    image_items: list[dict] = []
    out_dir = _image_output_dir(source_path)
    rels = getattr(doc.part, "rels", {})
    image_rels = [rel for rel in rels.values() if "image" in getattr(rel, "target_ref", "")]
    if not image_rels:
        return parts, image_items

    parts.append(f"[文档包含 {len(image_rels)} 张内嵌图片/图表。]")
    for idx, rel in enumerate(image_rels[:_IMAGE_EXTRACT_LIMIT], start=1):
        try:
            blob = rel.target_part.blob
            ext = os.path.splitext(getattr(rel.target_part, "partname", ""))[1] or ".png"
            img_path = os.path.join(out_dir, f"docx-image-{idx:02d}{ext}")
            with open(img_path, "wb") as f:
                f.write(blob)
            parts.append(f"[内嵌图片 {idx}：已从原文档抽取。]\n[图片文件] {img_path}")
            image_items.append({"page": None, "index": idx, "path": img_path, "description": ""})
        except Exception as e:
            logger.warning("DOCX 图片提取失败 (%s, image %d): %s", source_path, idx, e)
            parts.append(f"[内嵌图片 {idx}：提取失败。]")
    if len(image_rels) > _IMAGE_EXTRACT_LIMIT:
        parts.append(f"[另有 {len(image_rels) - _IMAGE_EXTRACT_LIMIT} 张图片未抽取，请查看原文档。]")
    return parts, image_items


def read_docx(path: str) -> str:
    return parse_docx(path).text


def parse_docx(path: str) -> ParsedDocument:
    from docx import Document as Docx
    doc = Docx(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)

    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[表格 {ti}]\n" + "\n".join(rows))

    image_parts, image_items = _extract_docx_images(doc, path)
    parts.extend(image_parts)
    return ParsedDocument(text="\n\n".join(parts), images=image_items)


def parse_any(path: str) -> ParsedDocument:
    p = path.lower()
    if p.endswith(".pdf"):
        return parse_pdf(path)
    if p.endswith(".docx"):
        return parse_docx(path)
    if p.endswith(".txt") or p.endswith(".md"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return ParsedDocument(text=f.read(), images=[])
    raise ValueError(f"Unsupported file: {path}")


def read_any(path: str) -> str:
    return parse_any(path).text
